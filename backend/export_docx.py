"""根据前端 JSON 生成 .docx（对齐 trpg_converter.html，并支持背景色块与插图）。"""
from __future__ import annotations

import io
from collections import defaultdict
from typing import Any

from docx import Document as create_document
from docx.document import Document as DocxDocument
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

COLORS_HEX = ["7eb8a4", "e07b5a", "9b84c4", "5a9fd4", "d4a84b", "c47a7a"]
DICE_RGB = RGBColor(0x4A, 0x8A, 0x4A)

ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}


def _norm_hex6(color: str) -> str:
    c = color.strip().lstrip("#")
    if len(c) >= 6:
        return c[:6].upper()
    return "5A4A3A"


def _mix_white(hex6: str, tint: float = 0.18) -> str:
    """将填充色冲淡，接近预览里半透明底效果。"""
    h = _norm_hex6(hex6)
    r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
    t = max(0.0, min(1.0, tint))

    def m(x: int) -> int:
        return int(x * t + 255 * (1 - t))

    return f"{m(r):02X}{m(g):02X}{m(b):02X}"


def _para_shading(paragraph, fill_hex: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    p_pr = paragraph._p.get_or_add_pPr()
    p_pr.append(shd)


def _cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def _title_border(p) -> None:
    p_bdr = OxmlElement("w:pBdr")
    for side in ("top", "bottom", "left", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), "8B6F4A")
        p_bdr.append(el)
    p._p.get_or_add_pPr().append(p_bdr)


def _set_cell_margins_dxa(cell, **kwargs: int) -> None:
    """为单元格设置内边距，单位 DXA。可选 top、left、bottom、right。"""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for side in ("top", "left", "bottom", "right"):
        if side not in kwargs:
            continue
        w = int(kwargs[side])
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(w))
        el.set(qn("w:type"), "dxa")
        tc_mar.append(el)
    tc_pr.append(tc_mar)


def _rgb_for_character(
    name: str,
    char_map: dict[str, int],
    char_colors: dict[str, str],
    palette: list[str],
) -> RGBColor:
    raw = str(char_colors.get(name, "")).strip().lstrip("#")
    if len(raw) >= 6:
        h6 = raw[:6].upper()
        try:
            return RGBColor(int(h6[0:2], 16), int(h6[2:4], 16), int(h6[4:6], 16))
        except ValueError:
            pass
    ci = int(char_map.get(name, 0)) % len(palette)
    nc = palette[ci]
    return RGBColor(int(nc[0:2], 16), int(nc[2:4], 16), int(nc[4:6], 16))


def _set_run_east_asia(run, name: str = "宋体") -> None:
    run.font.name = name
    try:
        r_pr = run._element.get_or_add_rPr()
        rfonts = r_pr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            r_pr.insert(0, rfonts)
        rfonts.set(qn("w:ascii"), name)
        rfonts.set(qn("w:hAnsi"), name)
        rfonts.set(qn("w:eastAsia"), name)
    except (AttributeError, TypeError):
        pass


def _speech_content_fw_paren_start(content: str) -> bool:
    first = ((content or "").split("\n")[0]).lstrip()
    return first.startswith("（")


def _apply_body_line_spacing(pf, mult: float) -> None:
    """正文多倍行距；1.0 或不填时保持 Word 默认，不写行距规则。"""
    try:
        m = max(1.0, min(3.0, float(mult)))
        if m <= 1.0 + 1e-6:
            return
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = m
    except (TypeError, ValueError, AttributeError):
        pass


def _line_shade_fill(line_index: int, blocks: list[dict[str, Any]]) -> str | None:
    for b in blocks:
        lo = min(int(b["start"]), int(b["end"]))
        hi = max(int(b["start"]), int(b["end"]))
        if lo <= line_index <= hi:
            return _mix_white(str(b.get("color", "#5a4a3a")))
    return None


def _add_title_paragraph(doc: DocxDocument, text: str, title_style: dict[str, Any]) -> None:
    fs_pt = float(title_style.get("fontSizePt", 14))
    align = ALIGN_MAP.get(str(title_style.get("align", "center")), WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    if bool(title_style.get("border", True)):
        _title_border(p)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(fs_pt)
    _set_run_east_asia(run, "宋体")
    run.font.color.rgb = RGBColor(0x2C, 0x1F, 0x0F)


def build_docx_bytes(payload: dict[str, Any]) -> bytes:
    title_style: dict[str, Any] = dict(payload.get("titleStyle") or {})
    titles_raw: list[dict[str, Any]] = list(payload.get("insertedTitles") or [])
    legacy = payload.get("title") or {}
    if str(legacy.get("text", "")).strip() and not titles_raw:
        titles_raw = [{"text": str(legacy["text"]).strip(), "afterIndex": -1}]
        for k in ("fontSizePt", "align", "border"):
            if k in legacy and k not in title_style:
                title_style[k] = legacy[k]

    lines: list[dict[str, Any]] = list(payload.get("lines") or [])
    char_map: dict[str, int] = dict(payload.get("charMap") or {})
    bg_blocks: list[dict[str, Any]] = list(payload.get("backgroundBlocks") or [])
    hidden_raw = payload.get("hiddenLineIndices") or []
    hidden_set: set[int] = set()
    for x in hidden_raw:
        try:
            hidden_set.add(int(x))
        except (TypeError, ValueError):
            pass
    appearance: dict[str, Any] = dict(payload.get("appearance") or {})
    body_font = str(appearance.get("bodyFont", "宋体"))
    speaker_font = str(appearance.get("speakerFont", "黑体"))
    char_colors: dict[str, str] = {str(k): str(v) for k, v in dict(appearance.get("charColors") or {}).items()}

    try:
        body_pt = float(appearance.get("bodyFontSizePt", 11))
    except (TypeError, ValueError):
        body_pt = 11.0
    body_pt = max(6.0, min(36.0, body_pt))
    dice_bot_pt = body_pt * 9.0 / 11.0
    dice_content_pt = body_pt * 10.0 / 11.0

    page_layout: dict[str, Any] = dict(payload.get("pageLayout") or {})
    margin_top = float(page_layout.get("marginTopMm", 25.4))
    margin_bottom = float(page_layout.get("marginBottomMm", 25.4))
    margin_left = float(page_layout.get("marginLeftMm", 25.4))
    margin_right = float(page_layout.get("marginRightMm", 25.4))
    line_spacing_mult = float(page_layout.get("lineSpacingMultiple", 1.0))
    parens_speech_right = bool(payload.get("parensSpeechRightAlign"))

    n_lines = len(lines)
    titles_by: dict[int, list[dict[str, Any]]] = defaultdict(list)
    for tit in titles_raw:
        idx = int(tit.get("afterIndex", -1))
        idx = max(-1, min(idx, n_lines))
        titles_by[idx].append(tit)

    doc = create_document()
    sec = doc.sections[0]
    sec.page_height = Mm(297)
    sec.page_width = Mm(210)
    sec.left_margin = Mm(max(5.0, min(60.0, margin_left)))
    sec.right_margin = Mm(max(5.0, min(60.0, margin_right)))
    sec.top_margin = Mm(max(5.0, min(60.0, margin_top)))
    sec.bottom_margin = Mm(max(5.0, min(60.0, margin_bottom)))

    TOTAL_W = 8504

    def _new_nil_tbl_borders() -> Any:
        el = OxmlElement("w:tblBorders")
        for tag in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = OxmlElement(f"w:{tag}")
            b.set(qn("w:val"), "nil")
            el.append(b)
        return el

    def flush_after(after_idx: int) -> None:
        for tit in titles_by.get(after_idx, []):
            ttxt = str(tit.get("text", "")).strip()
            if ttxt:
                _add_title_paragraph(doc, ttxt, title_style)

    flush_after(-1)

    for i, l in enumerate(lines):
        if i in hidden_set:
            flush_after(i)
            continue

        shade = _line_shade_fill(i, bg_blocks)
        t = l.get("type")

        if t == "narration":
            content = str(l.get("content", ""))
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Mm(3.5)
            _apply_body_line_spacing(p.paragraph_format, line_spacing_mult)
            if shade:
                _para_shading(p, shade)
            run = p.add_run(content)
            run.font.size = Pt(body_pt)
            _set_run_east_asia(run, body_font)
            run.font.color.rgb = RGBColor(0x6B, 0x5F, 0x52)
            flush_after(i)
            continue

        if t == "dice":
            name = str(l.get("name", ""))
            content = str(l.get("content", ""))
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Mm(3.5)
            _apply_body_line_spacing(p.paragraph_format, line_spacing_mult)
            if shade:
                _para_shading(p, shade)
            r1 = p.add_run(f"[{name}] ")
            r1.font.size = Pt(dice_bot_pt)
            _set_run_east_asia(r1, body_font)
            r1.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            r2 = p.add_run(content)
            r2.font.size = Pt(dice_content_pt)
            _set_run_east_asia(r2, body_font)
            r2.font.color.rgb = DICE_RGB
            flush_after(i)
            continue

        if t == "speech":
            name = str(l.get("name", ""))
            content = str(l.get("content", ""))
            rgb = _rgb_for_character(name, char_map, char_colors, COLORS_HEX)

            table = doc.add_table(rows=1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.autofit = False
            tbl = table._tbl
            tbl_pr = tbl.tblPr
            if tbl_pr is None:
                tbl_pr = OxmlElement("w:tblPr")
                tbl.insert(0, tbl_pr)
            tbl_pr.append(_new_nil_tbl_borders())
            w = OxmlElement("w:tblW")
            w.set(qn("w:w"), str(TOTAL_W))
            w.set(qn("w:type"), "dxa")
            tbl_pr.append(w)

            row = table.rows[0]
            c0 = row.cells[0]
            c1 = row.cells[1]
            if shade:
                _cell_shading(c0, shade)
                _cell_shading(c1, shade)

            _set_cell_margins_dxa(c0, right=280)
            _set_cell_margins_dxa(c1, left=100)

            c0.width = Mm(32)
            c1.width = Mm(126)
            c0.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            c1.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

            p0 = c0.paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p0.paragraph_format.space_before = Pt(2)
            p0.paragraph_format.space_after = Pt(2)
            _apply_body_line_spacing(p0.paragraph_format, line_spacing_mult)
            r0 = p0.add_run(name)
            r0.bold = True
            r0.font.size = Pt(body_pt)
            _set_run_east_asia(r0, speaker_font)
            r0.font.color.rgb = rgb

            parts = content.split("\n") if content else [""]
            content_right = parens_speech_right and _speech_content_fw_paren_start(content)
            for li, line_txt in enumerate(parts):
                p1 = c1.paragraphs[0] if li == 0 else c1.add_paragraph()
                p1.paragraph_format.space_before = Pt(2)
                p1.paragraph_format.space_after = Pt(2)
                _apply_body_line_spacing(p1.paragraph_format, line_spacing_mult)
                if content_right:
                    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                r1 = p1.add_run(line_txt)
                r1.font.size = Pt(body_pt)
                _set_run_east_asia(r1, body_font)
                r1.font.color.rgb = rgb

            flush_after(i)
            continue

        flush_after(i)

    flush_after(len(lines))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
